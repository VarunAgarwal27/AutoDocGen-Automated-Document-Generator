from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import markdown
from docx import Document


# PDF formating

from bs4 import BeautifulSoup

def export_pdf(project_name: str, docs: str, path: str):
    pdf = SimpleDocTemplate(path, pagesize=A4)

    styles = getSampleStyleSheet()
    story = []

    # Title
    story.append(Paragraph(project_name, styles["Title"]))
    story.append(Spacer(1, 20))

    # Convert markdown → HTML → parse
    html = markdown.markdown(docs)
    soup = BeautifulSoup(html, "html.parser")

    for element in soup:
        if element.name == "h1":
            story.append(Paragraph(f"<b>{element.text}</b>", styles["Heading1"]))
        elif element.name == "h2":
            story.append(Paragraph(f"<b>{element.text}</b>", styles["Heading2"]))
        elif element.name == "h3":
            story.append(Paragraph(f"<b>{element.text}</b>", styles["Heading3"]))
        elif element.name == "p":
            story.append(Paragraph(element.text, styles["Normal"]))
        elif element.name == "ul":
            for li in element.find_all("li"):
                story.append(Paragraph(f"• {li.text}", styles["Normal"]))

        story.append(Spacer(1, 12))

    pdf.build(story)


# Document formating
from bs4 import BeautifulSoup

def export_docx(project_name: str, docs: str, path: str):
    document = Document()

    document.add_heading(project_name, level=0)

    html = markdown.markdown(docs)
    soup = BeautifulSoup(html, "html.parser")

    for element in soup:
        if element.name == "h1":
            document.add_heading(element.text, level=1)
        elif element.name == "h2":
            document.add_heading(element.text, level=2)
        elif element.name == "h3":
            document.add_heading(element.text, level=3)
        elif element.name == "p":
            document.add_paragraph(element.text)
        elif element.name == "ul":
            for li in element.find_all("li"):
                document.add_paragraph(li.text, style="List Bullet")

    document.save(path)
